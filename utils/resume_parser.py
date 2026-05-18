"""
resume_parser.py
Extracts and structures all resume data from PDF or DOCX files.
Uses regex + NLP-style pattern matching with skill normalization,
section detection, and categorization — no heavy ML dependencies needed.
"""

import csv
import html
import io
import mimetypes
import os
import re
import zipfile
from difflib import SequenceMatcher

import pdfplumber
from docx import Document

# ── Skill Alias Normalization Map ─────────────────────────────────────────────
SKILL_ALIASES = {
    "js": "javascript", "ts": "typescript", "py": "python",
    "ml": "machine learning", "dl": "deep learning", "ai": "artificial intelligence",
    "nlp": "natural language processing", "cv": "computer vision",
    "k8s": "kubernetes", "tf": "tensorflow", "pt": "pytorch",
    "pg": "postgresql", "mongo": "mongodb", "es": "elasticsearch",
    "gql": "graphql", "node.js": "node", "nodejs": "node",
    "reactjs": "react", "react.js": "react", "vuejs": "vue", "vue.js": "vue",
    "angularjs": "angular", "next.js": "next", "nextjs": "next",
    "express.js": "express", "expressjs": "express",
    "sklearn": "scikit-learn", "scikit learn": "scikit-learn",
    "pandas": "pandas", "numpy": "numpy",
    "c plus plus": "c++", "cplusplus": "c++",
    "c sharp": "c#", "csharp": "c#",
    "dotnet": ".net", "dot net": ".net",
    "aws": "aws", "amazon web services": "aws",
    "gcp": "gcp", "google cloud": "gcp",
    "azure": "azure", "microsoft azure": "azure",
    "ci cd": "ci/cd", "cicd": "ci/cd",
    "rest": "rest api", "restful": "rest api", "restful api": "rest api",
    "sql server": "sql", "mysql": "mysql", "postgres": "postgresql",
    "powerbi": "power bi", "power-bi": "power bi",
    "tableau": "tableau", "looker": "looker",
    "git hub": "github", "gitlab": "gitlab",
    "linux": "linux", "unix": "linux",
    "oop": "object oriented programming", "oops": "object oriented programming",
    "dsa": "data structures", "ds": "data structures",
    "html5": "html", "css3": "css",
    "sass": "sass", "scss": "sass",
    "flutter": "flutter", "dart": "dart",
    "kotlin": "kotlin", "swift": "swift",
    "springboot": "spring boot", "spring-boot": "spring boot",
    "fastapi": "fastapi", "flask": "flask", "django": "django",
    "hadoop": "hadoop", "spark": "spark", "kafka": "kafka",
    "airflow": "airflow", "dbt": "dbt",
    "selenium": "selenium", "cypress": "cypress",
    "jest": "jest", "pytest": "pytest",
    "redis": "redis", "rabbitmq": "rabbitmq",
    "terraform": "terraform", "ansible": "ansible",
    "jenkins": "jenkins", "github actions": "github actions",
    "docker": "docker", "kubernetes": "kubernetes",
    "figma": "figma", "sketch": "sketch", "adobe xd": "adobe xd",
    "agile": "agile", "scrum": "scrum", "kanban": "kanban",
    "jira": "jira", "confluence": "confluence",
    "excel": "excel", "powerpoint": "powerpoint",
    "r language": "r", "rlang": "r",
    "matlab": "matlab", "julia": "julia",
    "rust": "rust", "go lang": "golang", "golang": "golang",
    "solidity": "solidity", "web3": "web3",
}

# ── Master Skill Taxonomy ─────────────────────────────────────────────────────
SKILL_TAXONOMY = {
    "Languages": [
        "python", "java", "javascript", "typescript", "c++", "c#", "c", "go",
        "golang", "rust", "kotlin", "swift", "dart", "r", "matlab", "scala",
        "php", "ruby", "perl", "bash", "shell", "powershell", "solidity",
        "html", "css", "sass", "sql",
    ],
    "Frontend": [
        "react", "angular", "vue", "next", "nuxt", "svelte", "redux",
        "tailwind", "bootstrap", "material ui", "chakra ui", "webpack",
        "vite", "figma", "sketch", "adobe xd", "graphql", "rest api",
    ],
    "Backend": [
        "node", "express", "django", "flask", "fastapi", "spring boot",
        "spring", "laravel", "rails", "asp.net", ".net", "fastify",
        "rest api", "graphql", "microservices", "grpc", "websocket",
    ],
    "Databases": [
        "mysql", "postgresql", "mongodb", "sqlite", "redis", "elasticsearch",
        "cassandra", "dynamodb", "firebase", "supabase", "oracle", "sql server",
        "neo4j", "influxdb", "cockroachdb",
    ],
    "AI/ML": [
        "machine learning", "deep learning", "natural language processing",
        "computer vision", "tensorflow", "pytorch", "scikit-learn", "keras",
        "hugging face", "transformers", "openai", "langchain", "llm",
        "pandas", "numpy", "matplotlib", "seaborn", "plotly", "scipy",
        "xgboost", "lightgbm", "mlops", "mlflow", "kubeflow",
    ],
    "DevOps/Cloud": [
        "docker", "kubernetes", "aws", "azure", "gcp", "terraform", "ansible",
        "jenkins", "github actions", "gitlab ci", "ci/cd", "linux", "nginx",
        "apache", "prometheus", "grafana", "elk stack", "helm", "istio",
    ],
    "Data Engineering": [
        "spark", "hadoop", "kafka", "airflow", "dbt", "data pipelines",
        "etl", "snowflake", "databricks", "bigquery", "redshift", "hive",
        "flink", "nifi", "tableau", "power bi", "looker",
    ],
    "Mobile": [
        "flutter", "react native", "android", "ios", "kotlin", "swift",
        "jetpack compose", "swiftui", "firebase", "android sdk",
    ],
    "Tools": [
        "git", "github", "gitlab", "jira", "confluence", "postman",
        "swagger", "vs code", "intellij", "eclipse", "xcode", "linux",
        "agile", "scrum", "kanban", "excel", "powerpoint",
    ],
    "Soft Skills": [
        "communication", "teamwork", "leadership", "problem solving",
        "critical thinking", "time management", "adaptability",
        "collaboration", "presentation", "mentoring", "project management",
    ],
}

# Flat lookup: skill → category
_SKILL_TO_CATEGORY = {
    skill: cat
    for cat, skills in SKILL_TAXONOMY.items()
    for skill in skills
}

# All known skills as a flat set for fast lookup
_ALL_KNOWN_SKILLS = set(_SKILL_TO_CATEGORY.keys())

# ── Section Header Patterns ───────────────────────────────────────────────────
SECTION_PATTERNS = {
    "skills":       re.compile(r"\b(technical\s+)?skills?\b", re.I),
    "experience":   re.compile(r"\b(work\s+)?experience\b|\bemployment\b|\bwork\s+history\b", re.I),
    "education":    re.compile(r"\beducation\b|\bacademic\b|\bqualification\b", re.I),
    "projects":     re.compile(r"\bprojects?\b|\bportfolio\b|\bwork\s+samples?\b", re.I),
    "certifications": re.compile(r"\bcertif\w+\b|\bcourses?\b|\btraining\b|\bachievements?\b", re.I),
    "summary":      re.compile(r"\b(professional\s+)?summary\b|\bobjective\b|\bprofile\b|\babout\b", re.I),
}

# ── Text Extraction ───────────────────────────────────────────────────────────

def extract_text_from_pdf(file_obj) -> str:
    """Extract all text from a PDF file object using pdfplumber."""
    lines = []
    with pdfplumber.open(file_obj) as pdf:
        for page in pdf.pages:
            # Try table extraction first (handles columnar resumes)
            tables = page.extract_tables()
            if tables:
                for table in tables:
                    for row in table:
                        lines.append(" | ".join(cell or "" for cell in row))
            text = page.extract_text(x_tolerance=3, y_tolerance=3)
            if text:
                lines.append(text)
    return "\n".join(lines)


def extract_text_from_docx(file_obj) -> str:
    """Extract all text from a DOCX file object."""
    doc = Document(file_obj)
    lines = []
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(lines)


def _read_bytes(file_obj) -> bytes:
    if hasattr(file_obj, "getvalue"):
        return file_obj.getvalue()
    current = file_obj.tell() if hasattr(file_obj, "tell") else None
    data = file_obj.read()
    if current is not None and hasattr(file_obj, "seek"):
        file_obj.seek(current)
    return data


def extract_text_from_plain_file(file_obj) -> str:
    """Extract text from common plain-text resume uploads."""
    data = _read_bytes(file_obj)
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="ignore")


def extract_text_from_csv(file_obj) -> str:
    """Flatten CSV/TSV rows into text that the resume parser can score."""
    text = extract_text_from_plain_file(file_obj)
    delimiter = "\t" if "\t" in text[:1000] else ","
    rows = csv.reader(io.StringIO(text), delimiter=delimiter)
    return "\n".join(" | ".join(cell.strip() for cell in row if cell.strip()) for row in rows)


def extract_text_from_rtf(file_obj) -> str:
    """Best-effort RTF text extraction without extra dependencies."""
    text = extract_text_from_plain_file(file_obj)
    text = re.sub(r"{\\\*[^{}]*}", " ", text)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+\d* ?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    return re.sub(r"\s+", " ", text).strip()


def extract_text_from_html(file_obj) -> str:
    """Extract visible-ish text from HTML resume exports."""
    text = extract_text_from_plain_file(file_obj)
    text = re.sub(r"(?is)<(script|style).*?>.*?</\1>", " ", text)
    text = re.sub(r"(?s)<br\s*/?>|</p>|</div>|</li>|</tr>", "\n", text)
    text = re.sub(r"(?s)<.*?>", " ", text)
    return html.unescape(re.sub(r"[ \t]+", " ", text)).strip()


def extract_text_from_odt(file_obj) -> str:
    """Extract text from OpenDocument text files."""
    data = _read_bytes(file_obj)
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        content = zf.read("content.xml").decode("utf-8", errors="ignore")
    content = re.sub(r"(?s)<text:line-break\s*/>|</text:p>|</text:h>", "\n", content)
    content = re.sub(r"(?s)<.*?>", " ", content)
    return html.unescape(re.sub(r"[ \t]+", " ", content)).strip()


def _extract_text_from_image_with_gemini(data: bytes, filename: str) -> str:
    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        return ""

    try:
        import google.generativeai as genai
    except ImportError:
        return ""

    mime_type = mimetypes.guess_type(filename)[0] or "image/png"
    genai.configure(api_key=api_key)
    model_name = os.environ.get("GEMINI_OCR_MODEL", "gemini-1.5-flash")
    model = genai.GenerativeModel(model_name)
    response = model.generate_content([
        "Extract only the readable resume text from this image. Preserve line breaks where useful.",
        {"mime_type": mime_type, "data": data},
    ])
    return (getattr(response, "text", "") or "").strip()


def _extract_text_from_image_with_tesseract(data: bytes) -> str:
    try:
        from PIL import Image
        import pytesseract
    except ImportError as exc:
        raise RuntimeError("Local OCR packages are unavailable.") from exc

    image = Image.open(io.BytesIO(data))
    return pytesseract.image_to_string(image).strip()


def _extract_text_from_image_with_easyocr(data: bytes) -> str:
    try:
        from PIL import Image
        import easyocr
        import numpy as np
    except ImportError as exc:
        raise RuntimeError("Bundled OCR packages are unavailable.") from exc

    languages = [
        lang.strip()
        for lang in os.environ.get("OCR_LANGUAGES", "en").split(",")
        if lang.strip()
    ]
    image = Image.open(io.BytesIO(data)).convert("RGB")
    reader = easyocr.Reader(languages or ["en"], gpu=False)
    lines = reader.readtext(np.array(image), detail=0, paragraph=True)
    return "\n".join(line.strip() for line in lines if line and line.strip())


def extract_text_from_image(file_obj, filename: str) -> str:
    """Extract text from image resumes using app-side OCR services."""
    data = _read_bytes(file_obj)
    errors = []

    try:
        text = _extract_text_from_image_with_gemini(data, filename)
        if text:
            return text
    except Exception as exc:
        errors.append(f"cloud OCR: {exc}")

    try:
        text = _extract_text_from_image_with_tesseract(data)
        if text:
            return text
    except Exception as exc:
        errors.append(f"local OCR: {exc}")

    try:
        text = _extract_text_from_image_with_easyocr(data)
        if text:
            return text
    except Exception as exc:
        errors.append(f"bundled OCR: {exc}")

    detail = " ".join(errors) if errors else "No OCR provider is configured."
    raise ValueError(
        "Image resume text could not be read automatically on this deployment. "
        "Please try uploading a clearer image, or use a PDF/DOCX/TXT version."
    ) from RuntimeError(detail)


def extract_text(file_obj, filename: str) -> str:
    """Route to correct extractor based on file extension."""
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return extract_text_from_pdf(file_obj)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(file_obj)
    elif ext in ("txt", "md", "markdown", "log"):
        return extract_text_from_plain_file(file_obj)
    elif ext in ("csv", "tsv"):
        return extract_text_from_csv(file_obj)
    elif ext == "rtf":
        return extract_text_from_rtf(file_obj)
    elif ext in ("html", "htm"):
        return extract_text_from_html(file_obj)
    elif ext == "odt":
        return extract_text_from_odt(file_obj)
    elif ext in ("jpg", "jpeg", "png", "webp", "bmp", "tif", "tiff"):
        return extract_text_from_image(file_obj, filename)
    raise ValueError(
        f"Unsupported file type: .{ext}. Try PDF, DOCX, TXT, CSV, RTF, HTML, ODT, JPG, or PNG."
    )


# ── Section Splitter ──────────────────────────────────────────────────────────

def split_into_sections(text: str) -> dict:
    """
    Split resume text into named sections.
    Returns dict: { section_name: [lines] }
    """
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    sections = {k: [] for k in SECTION_PATTERNS}
    sections["other"] = []
    current = "other"

    for line in lines:
        matched = False
        for sec, pattern in SECTION_PATTERNS.items():
            # Section headers are usually short lines (< 5 words)
            if pattern.search(line) and len(line.split()) <= 6:
                current = sec
                matched = True
                break
        if not matched:
            sections[current].append(line)

    return sections


# ── Skill Normalization ───────────────────────────────────────────────────────

def normalize_skill(raw: str) -> str:
    """Normalize a raw skill string using alias map and fuzzy matching."""
    cleaned = raw.strip().lower()
    cleaned = re.sub(r"[^\w\s\+\#\./]", "", cleaned)
    cleaned = re.sub(r"\s+", " ", cleaned).strip()

    # Direct alias lookup
    if cleaned in SKILL_ALIASES:
        return SKILL_ALIASES[cleaned]

    # Direct known skill
    if cleaned in _ALL_KNOWN_SKILLS:
        return cleaned

    # Fuzzy match against known skills (threshold 0.82)
    best, best_score = cleaned, 0.0
    for known in _ALL_KNOWN_SKILLS:
        score = SequenceMatcher(None, cleaned, known).ratio()
        if score > best_score:
            best_score = score
            best = known
    if best_score >= 0.82:
        return best

    return cleaned  # Return as-is if no match found


# ── Skill Extraction ──────────────────────────────────────────────────────────

def extract_skills_from_text(text: str) -> list[str]:
    """
    Extract skills using multi-pass strategy:
    1. Multi-word skill phrases (longest match first)
    2. Single-word tokens
    3. Comma/pipe/bullet separated lists
    """
    text_lower = text.lower()
    found = set()

    # Pass 1: match multi-word known skills first (longest first)
    sorted_skills = sorted(_ALL_KNOWN_SKILLS, key=len, reverse=True)
    for skill in sorted_skills:
        pattern = r"\b" + re.escape(skill) + r"\b"
        if re.search(pattern, text_lower):
            found.add(skill)

    # Pass 2: match aliases
    for alias, canonical in SKILL_ALIASES.items():
        pattern = r"\b" + re.escape(alias) + r"\b"
        if re.search(pattern, text_lower):
            found.add(canonical)

    # Pass 3: extract comma/pipe/bullet separated tokens and normalize
    token_pattern = re.compile(
        r"(?:^|[,|•·▪\-–—/\n])\s*([A-Za-z][A-Za-z0-9\s\+\#\.]{1,30}?)(?=[,|•·▪\-–—/\n]|$)",
        re.MULTILINE
    )
    for match in token_pattern.finditer(text):
        token = match.group(1).strip()
        if 2 <= len(token) <= 30:
            normalized = normalize_skill(token)
            if normalized in _ALL_KNOWN_SKILLS:
                found.add(normalized)

    return sorted(found)


# ── Contact Info Extraction ───────────────────────────────────────────────────

def extract_contact(text: str) -> dict:
    email   = re.search(r"[\w.+-]+@[\w-]+\.\w{2,}", text)
    phone   = re.search(r"(\+?\d[\d\s\-().]{8,15}\d)", text)
    linkedin = re.search(r"linkedin\.com/in/[\w\-]+", text, re.I)
    github  = re.search(r"github\.com/[\w\-]+", text, re.I)
    name_line = text.strip().splitlines()[0] if text.strip() else ""

    return {
        "name":     name_line[:60].strip() if len(name_line.split()) <= 6 else "",
        "email":    email.group(0) if email else "",
        "phone":    phone.group(1).strip() if phone else "",
        "linkedin": linkedin.group(0) if linkedin else "",
        "github":   github.group(0) if github else "",
    }


# ── Education Extraction ──────────────────────────────────────────────────────

def extract_education(lines: list[str]) -> list[dict]:
    """Parse education section into structured entries."""
    degree_pattern = re.compile(
        r"\b(b\.?tech|m\.?tech|b\.?e|m\.?e|b\.?sc|m\.?sc|bca|mca|"
        r"bachelor|master|phd|doctorate|diploma|b\.?com|mba|b\.?a|m\.?a)\b",
        re.I
    )
    year_pattern = re.compile(r"\b(19|20)\d{2}\b")
    entries, current = [], {}

    for line in lines:
        if degree_pattern.search(line):
            if current:
                entries.append(current)
            current = {"degree": line.strip(), "institution": "", "year": ""}
        elif current:
            if not current["institution"] and len(line.split()) >= 2:
                current["institution"] = line.strip()
            years = year_pattern.findall(line)
            if years and not current["year"]:
                current["year"] = " - ".join(years[-2:]) if len(years) >= 2 else years[-1]

    if current:
        entries.append(current)
    return entries


# ── Experience Extraction ─────────────────────────────────────────────────────

def extract_experience(lines: list[str]) -> list[dict]:
    """Parse experience section into structured job entries."""
    date_pattern = re.compile(
        r"\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec|january|february|"
        r"march|april|june|july|august|september|october|november|december)"
        r"[\s,]*\d{4}\b|\b\d{4}\s*[-–]\s*(\d{4}|present|current)\b",
        re.I
    )
    entries, current = [], {}

    for line in lines:
        if date_pattern.search(line) or (len(line.split()) <= 8 and any(
            w in line.lower() for w in ["engineer", "developer", "analyst", "intern",
                                         "manager", "lead", "architect", "scientist"]
        )):
            if current:
                entries.append(current)
            current = {
                "title": line.strip(),
                "company": "",
                "duration": date_pattern.search(line).group(0) if date_pattern.search(line) else "",
                "responsibilities": [],
            }
        elif current:
            if not current["company"] and len(line.split()) <= 6:
                current["company"] = line.strip()
            elif line.startswith(("•", "-", "–", "▪", "*")) or len(line) > 30:
                current["responsibilities"].append(line.lstrip("•-–▪* ").strip())

    if current:
        entries.append(current)
    return entries


# ── Project Extraction ────────────────────────────────────────────────────────

def extract_projects(lines: list[str]) -> list[dict]:
    """Parse projects section into structured entries."""
    entries, current = [], {}
    tech_pattern = re.compile(r"\b(tech(?:nologies)?|tools?|stack|built\s+with|using)\s*[:\-]?\s*(.+)", re.I)

    for line in lines:
        # Project title: short line, possibly bold/capitalized
        if len(line.split()) <= 8 and not line.startswith(("•", "-", "–")) and len(line) > 3:
            if current:
                entries.append(current)
            current = {"title": line.strip(), "description": "", "tech": []}
        elif current:
            tech_match = tech_pattern.search(line)
            if tech_match:
                raw_tech = re.split(r"[,|/]", tech_match.group(2))
                current["tech"] = [normalize_skill(t) for t in raw_tech if t.strip()]
            elif not current["description"]:
                current["description"] = line.strip()
            else:
                current["description"] += " " + line.strip()

    if current:
        entries.append(current)
    return entries


# ── Skill Categorization ──────────────────────────────────────────────────────

def categorize_skills(skills: list[str]) -> dict:
    """Group skills into taxonomy categories."""
    categorized = {cat: [] for cat in SKILL_TAXONOMY}
    categorized["Other"] = []

    for skill in skills:
        cat = _SKILL_TO_CATEGORY.get(skill.lower(), "Other")
        categorized[cat].append(skill)

    # Remove empty categories
    return {k: v for k, v in categorized.items() if v}


# ── Master Parse Function ─────────────────────────────────────────────────────

def parse_resume(file_obj, filename: str) -> dict:
    """
    Full resume parsing pipeline.

    Returns structured dict:
    {
        contact, skills, categorized_skills,
        education, experience, projects,
        certifications, summary, raw_text
    }
    """
    raw_text = extract_text(file_obj, filename)
    sections = split_into_sections(raw_text)

    # Extract skills from dedicated skills section + full text
    skills_text = "\n".join(sections.get("skills", []))
    all_skills_raw = extract_skills_from_text(raw_text)
    skills_from_section = extract_skills_from_text(skills_text)

    # Merge, deduplicate, normalize
    all_skills = sorted(set(
        normalize_skill(s) for s in (all_skills_raw + skills_from_section)
    ))

    # Separate technical vs soft skills
    soft_skill_set = set(SKILL_TAXONOMY.get("Soft Skills", []))
    technical_skills = [s for s in all_skills if s not in soft_skill_set]
    soft_skills      = [s for s in all_skills if s in soft_skill_set]

    return {
        "contact":            extract_contact(raw_text),
        "skills":             technical_skills,
        "soft_skills":        soft_skills,
        "categorized_skills": categorize_skills(technical_skills),
        "education":          extract_education(sections.get("education", [])),
        "experience":         extract_experience(sections.get("experience", [])),
        "projects":           extract_projects(sections.get("projects", [])),
        "certifications":     sections.get("certifications", []),
        "summary":            " ".join(sections.get("summary", []))[:500],
        "raw_text":           raw_text,
    }
